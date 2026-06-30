from __future__ import annotations

import math
import shutil
import zipfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SRC_DIR = Path(r"C:\Users\marce\AppData\Local\Temp")
IMG_DIR = ROOT / "prints_usuario"
ANN_DIR = ROOT / "prints_anotados"
OUT_DOCX = ROOT / "Manual_de_Uso_CafeFinance.docx"

COLORS = {
    "brown": RGBColor(62, 25, 10),
    "orange": RGBColor(203, 111, 28),
    "cream": RGBColor(255, 246, 226),
    "muted": RGBColor(92, 65, 45),
    "red": (226, 0, 0),
}


SCREENS = [
    {
        "key": "login",
        "title": "1. Entrar no sistema",
        "src": "codex-clipboard-f4512df5-38a2-4f3f-907a-9f974d4b44ae.png",
        "caption": "Tela de login: use e-mail e senha cadastrados para acessar o sistema.",
        "steps": [
            "Digite o e-mail cadastrado.",
            "Digite a senha.",
            "Marque Lembrar acesso apenas se estiver usando um computador seguro.",
            "Clique em Entrar.",
            "Caso ainda não tenha conta, clique em Criar conta.",
        ],
        "marks": [
            ("rect", (0.54, 0.37, 0.86, 0.45), "E-mail"),
            ("rect", (0.54, 0.50, 0.86, 0.58), "Senha"),
            ("rect", (0.54, 0.64, 0.86, 0.72), "Entrar"),
            ("rect", (0.63, 0.82, 0.76, 0.88), "Criar conta"),
        ],
    },
    {
        "key": "cadastro",
        "title": "2. Criar uma conta",
        "src": "codex-clipboard-3f485b8f-cc71-4dd0-adb4-b3322f041045.png",
        "caption": "Tela de cadastro: crie seu acesso informando nome, e-mail e senha.",
        "steps": [
            "Informe seu nome.",
            "Informe um e-mail válido.",
            "Digite a senha e confirme a mesma senha no campo ao lado.",
            "Clique em Criar conta para finalizar o cadastro.",
            "Se já tiver cadastro, clique em Entrar.",
        ],
        "marks": [
            ("rect", (0.28, 0.48, 0.49, 0.55), "Nome"),
            ("rect", (0.50, 0.48, 0.72, 0.55), "E-mail"),
            ("rect", (0.28, 0.61, 0.49, 0.67), "Senha"),
            ("rect", (0.50, 0.61, 0.72, 0.67), "Confirmar senha"),
            ("rect", (0.28, 0.69, 0.72, 0.77), "Criar conta"),
        ],
    },
    {
        "key": "perfil",
        "title": "3. Perfil - menu principal",
        "src": "codex-clipboard-58130d46-e23e-4a94-ae09-01b49695a346.png",
        "caption": "Perfil: menu principal aberto após o login, com o resumo financeiro da pessoa logada.",
        "steps": [
            "Confira o saldo atual, receitas e despesas.",
            "Veja as últimas movimentações registradas.",
            "Use + Novo registro para ir rapidamente para a tela de movimentações.",
            "Acompanhe os gastos por categoria.",
            "Use o menu superior para navegar entre Perfil, Movimentações, Relatórios, Economias e Sair.",
        ],
        "marks": [
            ("rect", (0.02, 0.28, 0.33, 0.45), "Saldo"),
            ("rect", (0.34, 0.28, 0.65, 0.45), "Receitas"),
            ("rect", (0.66, 0.28, 0.97, 0.45), "Despesas"),
            ("rect", (0.46, 0.52, 0.57, 0.59), "+ Novo registro"),
            ("rect", (0.59, 0.48, 0.98, 0.66), "Categorias"),
            ("ellipse", (0.41, 0.02, 0.93, 0.11), "Menu"),
        ],
    },
    {
        "key": "movimentacoes_novo",
        "title": "4. Registrar entradas e saídas",
        "src": "codex-clipboard-78c05f5c-9195-46a1-873e-32560e4ffb93.png",
        "caption": "Movimentações: registre receitas, despesas e compras parceladas.",
        "steps": [
            "Escolha Entrada ou Saída.",
            "Digite o valor. O sistema formata em reais automaticamente.",
            "Selecione a data pelo calendário.",
            "Escolha a categoria da movimentação.",
            "Use a descrição opcional para identificar melhor o registro.",
            "Clique em Salvar entrada ou Salvar saída.",
        ],
        "marks": [
            ("rect", (0.04, 0.08, 0.41, 0.15), "Entrada/Saída"),
            ("rect", (0.04, 0.22, 0.22, 0.28), "Valor"),
            ("rect", (0.22, 0.22, 0.41, 0.28), "Data"),
            ("rect", (0.03, 0.31, 0.41, 0.54), "Categorias"),
            ("rect", (0.03, 0.64, 0.41, 0.77), "Descrição"),
            ("rect", (0.03, 0.94, 0.41, 1.00), "Salvar"),
        ],
    },
    {
        "key": "movimentacoes_historico",
        "title": "5. Consultar, editar e excluir registros",
        "src": "codex-clipboard-677097fe-a705-4ad9-81b4-460e586b421d.png",
        "caption": "Histórico: os registros ficam agrupados por mês e podem ser editados ou excluídos.",
        "steps": [
            "Clique no mês para abrir os registros daquele período.",
            "Confira entradas, saídas e saldo do mês.",
            "Use Editar para alterar um registro.",
            "Use Excluir para remover um registro.",
            "A lista interna possui rolagem quando houver muitos registros.",
        ],
        "marks": [
            ("rect", (0.44, 0.13, 0.57, 0.36), "Meses"),
            ("rect", (0.44, 0.39, 0.73, 0.45), "Resumo do mês"),
            ("rect", (0.74, 0.48, 0.79, 0.55), "Editar"),
            ("rect", (0.79, 0.48, 0.85, 0.55), "Excluir"),
            ("rect", (0.93, 0.45, 0.96, 0.93), "Rolagem"),
        ],
    },
    {
        "key": "relatorios_resumo",
        "title": "6. Relatórios - dashboards do mês",
        "src": "codex-clipboard-d5f3300a-2de5-4fac-9709-176953d0137f.png",
        "caption": "Relatórios: área dos dashboards do sistema, com saldo, entradas, saídas e fatura parcelada por mês.",
        "steps": [
            "Selecione o mês desejado.",
            "Confira o saldo acumulado.",
            "Compare entradas e saídas do mês.",
            "Veja o total previsto de fatura parcelada.",
            "Use Registrar movimentação para voltar ao cadastro de registros.",
        ],
        "marks": [
            ("rect", (0.02, 0.10, 0.49, 0.20), "Meses"),
            ("rect", (0.02, 0.22, 0.25, 0.37), "Saldo"),
            ("rect", (0.26, 0.22, 0.49, 0.37), "Entradas"),
            ("rect", (0.50, 0.22, 0.73, 0.37), "Saídas"),
            ("rect", (0.74, 0.22, 0.97, 0.37), "Fatura"),
            ("rect", (0.81, 0.00, 0.97, 0.07), "Registrar"),
        ],
    },
    {
        "key": "relatorios_graficos",
        "title": "7. Dashboards de análise",
        "src": "codex-clipboard-c4318afb-e207-439f-9df4-164aa6801fb8.png",
        "caption": "Dashboards: acompanhe gastos por categoria e evolução do saldo no mês dentro da aba Relatórios.",
        "steps": [
            "No gráfico de categorias, veja onde o dinheiro está sendo mais gasto.",
            "A lista abaixo detalha o valor por categoria.",
            "No gráfico de evolução, compare entradas, saídas e saldo acumulado.",
        ],
        "marks": [
            ("ellipse", (0.23, 0.11, 0.40, 0.42), "Categorias"),
            ("rect", (0.03, 0.49, 0.56, 0.96), "Lista"),
            ("rect", (0.60, 0.05, 0.96, 0.50), "Evolução"),
        ],
    },
    {
        "key": "relatorios_fatura",
        "title": "8. Fatura parcelada prevista",
        "src": "codex-clipboard-e0211220-87d2-40f0-8ee8-fc4f39398be2.png",
        "caption": "Fatura parcelada: mostra quanto já está comprometido em cada mês.",
        "steps": [
            "Use este bloco para visualizar compras parceladas futuras.",
            "Cada barra representa o valor previsto para um mês.",
            "Abaixo do gráfico, o sistema lista o valor de cada mês.",
        ],
        "marks": [
            ("rect", (0.05, 0.20, 0.95, 0.62), "Gráfico"),
            ("rect", (0.04, 0.70, 0.49, 0.93), "Junho"),
            ("rect", (0.50, 0.70, 0.96, 0.93), "Julho"),
        ],
    },
    {
        "key": "economias_topo",
        "title": "9. Criar metas e guardar economia",
        "src": "codex-clipboard-32ae7cd3-59a3-42d5-b8dc-ed2387867cc4.png",
        "caption": "Economias: crie metas e registre valores guardados para acompanhar o progresso.",
        "steps": [
            "Em Criar meta, informe nome, valor da meta e data limite, se houver.",
            "Clique em Criar meta.",
            "Em Guardar economia, escolha uma meta já cadastrada.",
            "Informe o valor guardado, data e descrição opcional.",
            "Clique em Guardar economia.",
        ],
        "marks": [
            ("rect", (0.03, 0.42, 0.49, 0.94), "Criar meta"),
            ("rect", (0.52, 0.42, 0.96, 0.94), "Guardar economia"),
            ("rect", (0.03, 0.13, 0.97, 0.39), "Resumo"),
        ],
    },
    {
        "key": "economias_metas",
        "title": "10. Metas, cafés concluídos e histórico",
        "src": "codex-clipboard-68dcc43f-dbc5-4f2a-8989-0995e5109060.png",
        "caption": "Acompanhamento: veja metas ativas, metas concluídas e últimas economias registradas.",
        "steps": [
            "Em Metas cadastradas, acompanhe a porcentagem de cada meta.",
            "Use Editar para alterar a meta.",
            "Use Excluir para remover a meta.",
            "Em Cafés Concluídos, abra as metas já finalizadas.",
            "Em Últimas economias, edite ou exclua valores guardados.",
        ],
        "marks": [
            ("rect", (0.02, 0.08, 0.49, 0.74), "Metas"),
            ("rect", (0.51, 0.08, 0.97, 0.33), "Concluídos"),
            ("rect", (0.02, 0.78, 0.49, 1.00), "Histórico"),
            ("rect", (0.36, 0.35, 0.47, 0.43), "Editar/Excluir"),
        ],
    },
]


def ensure_dirs() -> None:
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    ANN_DIR.mkdir(parents=True, exist_ok=True)


def copy_sources() -> None:
    for screen in SCREENS:
        src = SRC_DIR / screen["src"]
        dst = IMG_DIR / f"{screen['key']}.png"
        if not src.exists():
            raise FileNotFoundError(src)
        shutil.copyfile(src, dst)


def font(size: int, bold: bool = True) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def rel_box(box: tuple[float, float, float, float], width: int, height: int) -> tuple[int, int, int, int]:
    return (
        int(box[0] * width),
        int(box[1] * height),
        int(box[2] * width),
        int(box[3] * height),
    )


def draw_label(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str) -> None:
    label_font = font(24)
    pad_x, pad_y = 12, 7
    bbox = draw.textbbox((0, 0), text, font=label_font)
    w = bbox[2] - bbox[0] + pad_x * 2
    h = bbox[3] - bbox[1] + pad_y * 2
    x, y = xy
    draw.rounded_rectangle((x, y, x + w, y + h), radius=12, fill=(226, 0, 0), outline=(255, 255, 255), width=2)
    draw.text((x + pad_x, y + pad_y - 2), text, fill=(255, 255, 255), font=label_font)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    red = COLORS["red"]
    draw.line((start, end), fill=red, width=8)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 24
    spread = math.pi / 7
    left = (end[0] - length * math.cos(angle - spread), end[1] - length * math.sin(angle - spread))
    right = (end[0] - length * math.cos(angle + spread), end[1] - length * math.sin(angle + spread))
    draw.polygon([end, left, right], fill=red)


def annotate_images() -> None:
    for screen in SCREENS:
        img = Image.open(IMG_DIR / f"{screen['key']}.png").convert("RGBA")
        overlay = Image.new("RGBA", img.size, (255, 255, 255, 0))
        draw = ImageDraw.Draw(overlay)
        w, h = img.size

        for kind, rel, label in screen["marks"]:
            box = rel_box(rel, w, h)
            if kind == "ellipse":
                draw.ellipse(box, outline=COLORS["red"], width=8)
            else:
                draw.rounded_rectangle(box, radius=16, outline=COLORS["red"], width=8)

            label_x = max(8, min(box[0], w - 220))
            label_y = max(8, box[1] - 46)
            draw_label(draw, (label_x, label_y), label)

            start = (label_x + 24, label_y + 42)
            end = ((box[0] + box[2]) // 2, (box[1] + box[3]) // 2)
            draw_arrow(draw, start, end)

        annotated = Image.alpha_composite(img, overlay).convert("RGB")
        annotated.save(ANN_DIR / f"{screen['key']}_anotado.png", quality=92)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "E5CDA9") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run(run, size: int, color: RGBColor, bold: bool = False) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def add_title(doc: Document, text: str, size: int = 30) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run(run, size, COLORS["brown"], True)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(text)
    set_run(run, 14, COLORS["muted"], False)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, 20, COLORS["brown"], True)


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF2D8")
    set_cell_border(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run(run, 10, COLORS["muted"], True)


def add_steps(doc: Document, steps: list[str]) -> None:
    for index, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{index}. ")
        set_run(r1, 10, COLORS["orange"], True)
        r2 = p.add_run(step)
        set_run(r2, 10, COLORS["brown"], False)


def add_image(doc: Document, image_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(9.2))
    p.paragraph_format.space_after = Pt(4)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    cap_run = cap.add_run(caption)
    set_run(cap_run, 9, COLORS["muted"], False)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    add_title(doc, "Manual de Uso - CafeFinance", 34)
    add_subtitle(doc, "Guia rápido para acessar, registrar movimentações, acompanhar dashboards em Relatórios e controlar economias.")
    add_note(
        doc,
        "As marcacoes em vermelho indicam onde clicar ou quais campos observar em cada tela. "
        "Use este manual como apoio durante a apresentacao e para orientar novos usuarios.",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Fluxo principal de uso")
    set_run(run, 16, COLORS["brown"], True)
    add_steps(
        doc,
        [
            "Criar conta ou entrar no sistema.",
            "Conferir o resumo financeiro no Perfil, que é o menu principal.",
            "Registrar entradas e saídas em Movimentações.",
            "Acompanhar dashboards e gráficos na aba Relatórios.",
            "Criar metas e guardar valores em Economias.",
        ],
    )

    for screen in SCREENS:
        doc.add_page_break()
        add_heading(doc, screen["title"])
        add_steps(doc, screen["steps"])
        add_image(doc, ANN_DIR / f"{screen['key']}_anotado.png", screen["caption"])

    doc.add_page_break()
    add_heading(doc, "Resumo final")
    add_note(
        doc,
        "O CafeFinance foi pensado para organizar a gestão individual: entradas, saídas, parcelas, dashboards em Relatórios e metas de economia ficam reunidos em um único sistema.",
    )
    add_steps(
        doc,
        [
            "Perfil: menu principal com resumo rápido do saldo, receitas, despesas e últimos registros.",
            "Movimentações: cadastro, histórico, edição e exclusão de entradas e saídas.",
            "Relatórios: dashboards por mês, categoria, evolução e fatura parcelada.",
            "Economias: criação de metas, valores guardados, progresso e cafés concluídos.",
            "Sair: encerra a sessão do usuário.",
        ],
    )

    doc.save(OUT_DOCX)


def scrub_docx_metadata(path: Path) -> None:
    # Remove os metadados pessoais mais comuns do pacote DOCX.
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "docProps/core.xml":
                text = data.decode("utf-8")
                text = text.replace("<dc:creator>python-docx</dc:creator>", "<dc:creator>CafeFinance</dc:creator>")
                text = text.replace("<cp:lastModifiedBy></cp:lastModifiedBy>", "<cp:lastModifiedBy>CafeFinance</cp:lastModifiedBy>")
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


def main() -> None:
    ensure_dirs()
    copy_sources()
    annotate_images()
    build_docx()
    scrub_docx_metadata(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
